"""
ProofPilot branded .docx generator — v3

Matches the Steadfast reference document exactly:

COVER LAYOUT
  Line 1: "SEO MARKET OPPORTUNITY"   — Bebas Neue 24pt Electric Blue, centered
  Line 2: "& COMPETITIVE ANALYSIS"   — Bebas Neue 38pt Dark Blue, centered
  Subtitle: italic gray 13pt, centered
  Info table: light-gray label col (2500 dxa) + white value col (6860 dxa), no header row
  Reveals callout box: Dark Blue bg, Neon Green Bebas header, white body bullets
  [COVER_END] marker in content → inserts page break

TYPOGRAPHY
  H1  — Bebas Neue 28pt Dark Blue
  H2  — Bebas Neue 18pt, alternates Dark Blue / Electric Blue per section
  H3  — Bebas Neue 13pt Electric Blue
  Body — Calibri 11pt
  All fonts: 4-axis (ascii, hAnsi, cs, eastAsia) for cross-platform safety

TABLES
  Brand tables: colored header row (alternates per H2 section), alternating body rows
  Info table: label/value, no header row, light-gray left col, gray borders
  Callout boxes: Dark Blue bg, Neon Green bold header, white body, Neon Green bullets

DOCUMENT STRUCTURE
  Section header (repeats every page): "PROOFPILOT | TITLE" right-aligned + blue rule
  Section footer (repeats every page): centered client + date
  Font embedding: BebasNeue-regular.ttf injected into DOCX zip post-save
  docDefaults: default spacing stripped (python-docx adds spacing after=200 by default)
"""

import os
import re
import zipfile
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colors ───────────────────────────────────────────
DARK_BLUE  = RGBColor(0x00, 0x18, 0x4D)   # #00184D
ELEC_BLUE  = RGBColor(0x00, 0x51, 0xFF)   # #0051FF
NEON_GREEN = RGBColor(0xC8, 0xFF, 0x00)   # #C8FF00
MID_GRAY   = RGBColor(0x66, 0x66, 0x66)   # #666666
LIGHT_GRAY = RGBColor(0xAA, 0xAA, 0xBB)   # footer / dividers
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_ROW  = RGBColor(0xF4, 0xF4, 0xF4)   # even body rows (also info-table label fill)

BODY_FONT    = "Calibri"
DISPLAY_FONT = "Bebas Neue"

FONTS_DIR      = Path(__file__).parent / "fonts"
BEBAS_NEUE_TTF = FONTS_DIR / "BebasNeue-regular.ttf"
TEMP_DIR = Path(os.environ.get("DOCS_DIR", str(Path(__file__).parent.parent / "temp_docs")))

# Prefixes that mark workflow progress lines in SSE — skip in DOCX
_STATUS_PREFIXES = (
    "Pulling", "Fetching", "Researching", "Building", "Loading",
    "Analyzing", "Computing", "Gathering", "Checking",
)


# ═══════════════════════════════════════════════════════════
# Entry point
# ═══════════════════════════════════════════════════════════

def generate_docx(job_id: str, job_data: dict) -> Path:
    """Generate a branded ProofPilot .docx and return its path."""
    TEMP_DIR.mkdir(exist_ok=True)

    content        = job_data["content"]
    client_name    = job_data["client_name"]
    workflow_title = job_data["workflow_title"]

    doc = Document()
    _set_margins(doc)
    _fix_document_defaults(doc)          # strip default spacing python-docx adds
    _setup_repeating_header(doc, workflow_title)  # section header (all pages)
    _setup_repeating_footer(doc, client_name)      # section footer (all pages)
    _render_markdown(doc, content)

    out_path = TEMP_DIR / f"{job_id}.docx"
    doc.save(out_path)
    _embed_fonts(out_path)               # inject BebasNeue TTF into zip
    return out_path


# ═══════════════════════════════════════════════════════════
# Document-level setup
# ═══════════════════════════════════════════════════════════

def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)


def _fix_document_defaults(doc: Document) -> None:
    """
    Strip the default paragraph spacing python-docx writes into docDefaults
    (spacing after=200, line=276). The reference has an empty <pPr/> with no
    spacing — so every paragraph is tight by default.

    Also replace the theme-resolved rPrDefault font with explicit Calibri on
    all 4 axes, so body text doesn't fall back to Cambria.
    """
    styles_el = doc.styles.element

    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        return

    # Remove default paragraph spacing
    pPrDefault = docDefaults.find(".//" + qn("w:pPrDefault"))
    if pPrDefault is not None:
        pPr = pPrDefault.find(qn("w:pPr"))
        if pPr is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is not None:
                pPr.remove(spacing)

    # Fix default font: explicit Calibri, not theme-resolved
    rPrDefault = docDefaults.find(".//" + qn("w:rPrDefault"))
    if rPrDefault is not None:
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                # Drop any *Theme attributes (e.g. w:asciiTheme, w:hAnsiTheme)
                for attr in list(rFonts.attrib.keys()):
                    if "heme" in attr:
                        del rFonts.attrib[attr]
                for ax in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                    rFonts.set(qn(ax), BODY_FONT)


def _set_font_all_axes(run, font_name: str) -> None:
    """Set font on all 4 Word axes so it renders correctly in every locale."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for ax in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(ax), font_name)
    run.font.name = font_name


# ═══════════════════════════════════════════════════════════
# Section header / footer (repeat on every page)
# ═══════════════════════════════════════════════════════════

def _setup_repeating_header(doc: Document, workflow_title: str) -> None:
    """
    PROOFPILOT  |  WORKFLOW TITLE  right-aligned
    ─────────────────────────────────  (Electric Blue rule)
    These live in the section header XML so they appear on every page.
    """
    section = doc.sections[0]
    header  = section.header

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Clear existing runs
    for r in list(p.runs):
        p._p.remove(r._r)

    brand = p.add_run("PROOFPILOT")
    brand.bold = True
    brand.font.size = Pt(10)
    brand.font.color.rgb = DARK_BLUE
    _set_font_all_axes(brand, DISPLAY_FONT)

    sep = p.add_run("  |  ")
    sep.font.size = Pt(9)
    sep.font.color.rgb = ELEC_BLUE
    _set_font_all_axes(sep, BODY_FONT)

    wf = p.add_run(workflow_title.upper())
    wf.font.size = Pt(9)
    wf.font.color.rgb = ELEC_BLUE
    _set_font_all_axes(wf, BODY_FONT)

    div = header.add_paragraph()
    d = div.add_run("─" * 90)
    d.font.size = Pt(7)
    d.font.color.rgb = ELEC_BLUE
    _set_font_all_axes(d, BODY_FONT)
    div.space_after = Pt(0)


def _setup_repeating_footer(doc: Document, client_name: str) -> None:
    """Centered footer: ProofPilot · client · date. Lives in section footer XML."""
    section = doc.sections[0]
    footer  = section.footer

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        p._p.remove(r._r)

    f = p.add_run(
        f"ProofPilot  ·  {client_name}  ·  {datetime.now().strftime('%B %d, %Y')}"
    )
    f.font.size = Pt(8)
    f.font.color.rgb = LIGHT_GRAY
    _set_font_all_axes(f, BODY_FONT)


# ═══════════════════════════════════════════════════════════
# Font embedding
# ═══════════════════════════════════════════════════════════

def _embed_fonts(docx_path: Path) -> None:
    """
    Post-process: inject BebasNeue-regular.ttf into the DOCX zip so the font
    renders correctly on any machine, even without Bebas Neue installed.
    Updates fontTable.xml and fontTable.xml.rels accordingly.
    """
    if not BEBAS_NEUE_TTF.exists():
        return  # Font file not available — skip silently

    font_bytes = BEBAS_NEUE_TTF.read_bytes()
    tmp = docx_path.with_suffix(".tmp.docx")

    FONT_DECL = (
        '<w:font w:name="Bebas Neue">'
        '<w:embedRegular'
        ' w:fontKey="{00000000-0000-0000-0000-000000000000}"'
        ' r:id="rId10" w:subsetted="0"/>'
        "</w:font>"
    )
    FONT_REL = (
        '<Relationship Id="rId10"'
        ' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"'
        ' Target="fonts/BebasNeue-regular.ttf"/>'
    )
    FONT_RELS_NEW = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + FONT_REL
        + "</Relationships>"
    )

    with zipfile.ZipFile(docx_path, "r") as zin:
        existing = {item.filename for item in zin.infolist()}

        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == "word/fontTable.xml":
                    text = data.decode("utf-8")
                    if "Bebas Neue" not in text:
                        text = text.replace("</w:fonts>", FONT_DECL + "</w:fonts>")
                    data = text.encode("utf-8")

                elif item.filename == "word/_rels/fontTable.xml.rels":
                    text = data.decode("utf-8")
                    if "rId10" not in text:
                        text = text.replace("</Relationships>", FONT_REL + "</Relationships>")
                    data = text.encode("utf-8")

                zout.writestr(item, data)

            # Embed font binary
            if "word/fonts/BebasNeue-regular.ttf" not in existing:
                zout.writestr("word/fonts/BebasNeue-regular.ttf", font_bytes)

            # Create fontTable.xml.rels if the generated DOCX didn't have one
            if "word/_rels/fontTable.xml.rels" not in existing:
                zout.writestr(
                    "word/_rels/fontTable.xml.rels",
                    FONT_RELS_NEW.encode("utf-8"),
                )

    tmp.replace(docx_path)


# ═══════════════════════════════════════════════════════════
# Cell-level XML helpers
# ═══════════════════════════════════════════════════════════

def _set_cell_background(cell, rgb: RGBColor) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex: str = "CCCCCC", sz: int = 4) -> None:
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tcPr.append(borders)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


# ═══════════════════════════════════════════════════════════
# Table builders
# ═══════════════════════════════════════════════════════════

def _parse_table_lines(lines: list) -> tuple:
    """Parse markdown pipe-table lines → (headers, rows)."""
    headers: list = []
    rows:    list = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        if re.match(r"^\|[-:\s|]+\|$", stripped):
            continue
        cells = [c.strip() for c in stripped.split("|")]
        cells = [c for c in cells if c != "" or len(cells) > 2]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if not cells:
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def _add_info_table(doc: Document, rows: list) -> None:
    """
    Company info / metadata table (the | | | two-column pattern).
    Matches reference exactly:
      - Label col: 2500 dxa, light-gray (#F4F4F4) fill, bold Calibri 11pt
      - Value col: 6860 dxa, white fill, regular Calibri 11pt
      - All borders: #CCCCCC sz=4
      - No header row
    """
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        tbl_row = table.rows[row_idx]

        label_cell = tbl_row.cells[0]
        value_cell = tbl_row.cells[1]

        _set_cell_background(label_cell, LIGHT_ROW)
        _set_cell_background(value_cell, WHITE)
        _set_cell_borders(label_cell, "CCCCCC", 4)
        _set_cell_borders(value_cell, "CCCCCC", 4)
        _set_cell_width(label_cell, 2500)
        _set_cell_width(value_cell, 6860)

        # Label (strip ** markdown bold)
        lp = label_cell.paragraphs[0]
        lp.space_before = Pt(3)
        lp.space_after  = Pt(3)
        label_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", row_data[0] if row_data else "").strip()
        lr = lp.add_run(label_text)
        lr.bold = True
        lr.font.size = Pt(11)
        _set_font_all_axes(lr, BODY_FONT)

        # Value
        vp = value_cell.paragraphs[0]
        vp.space_before = Pt(3)
        vp.space_after  = Pt(3)
        value_text = (row_data[1].strip() if len(row_data) > 1 else "")
        vr = vp.add_run(value_text)
        vr.font.size = Pt(11)
        _set_font_all_axes(vr, BODY_FONT)

    spacer = doc.add_paragraph()
    spacer.space_after = Pt(8)


def _add_brand_table(
    doc: Document,
    headers: list,
    rows: list,
    header_color: RGBColor = None,
) -> None:
    """
    Standard branded data table.
    Header row color alternates per H2 section (caller passes header_color).
    Body rows alternate white / light gray.
    """
    if not headers:
        return
    if header_color is None:
        header_color = DARK_BLUE

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_background(cell, header_color)
        p = cell.paragraphs[0]
        for run in p.runs:
            run.text = ""
        clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", header_text).strip()
        run = p.add_run(clean)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        _set_font_all_axes(run, BODY_FONT)
        p.space_before = Pt(3)
        p.space_after  = Pt(3)

    # Body rows
    for row_idx, row_data in enumerate(rows):
        tbl_row = table.rows[row_idx + 1]
        bg = LIGHT_ROW if row_idx % 2 == 1 else WHITE
        for col_idx in range(len(headers)):
            cell = tbl_row.cells[col_idx]
            _set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.text = ""
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            _inline_format_table(p, cell_text)
            p.space_before = Pt(2)
            p.space_after  = Pt(2)

    spacer = doc.add_paragraph()
    spacer.space_after = Pt(6)


def _inline_format_table(paragraph, text: str) -> None:
    """**bold**, *italic*, plain — 10pt Calibri, all 4 font axes."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10)
            _set_font_all_axes(r, BODY_FONT)
        else:
            for sub in re.split(r"(\*[^*]+\*)", part):
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    r = paragraph.add_run(sub[1:-1])
                    r.italic = True
                    r.font.size = Pt(10)
                    _set_font_all_axes(r, BODY_FONT)
                else:
                    r = paragraph.add_run(sub)
                    r.font.size = Pt(10)
                    _set_font_all_axes(r, BODY_FONT)


# ═══════════════════════════════════════════════════════════
# Callout box
# ═══════════════════════════════════════════════════════════

def _add_callout_box(doc: Document, lines: list) -> None:
    """
    Dark Blue (#00184D) callout box from >-prefixed blockquote lines.
      **BOLD HEADER** → Neon Green, 13pt Bebas Neue
      - bullet item   → Neon Green bullet dot, white body text
      regular text    → white 10pt, **bold** highlights in Neon Green
    """
    content_lines = [l for l in lines if l.strip()]
    if not content_lines:
        return

    pre = doc.add_paragraph()
    pre.space_before = Pt(4)
    pre.space_after  = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _set_cell_background(cell, DARK_BLUE)
    _set_cell_width(cell, 9360)

    first_para = True
    for line in content_lines:
        raw = line.strip()
        if not raw:
            continue

        p = cell.paragraphs[0] if first_para else cell.add_paragraph()
        first_para = False

        # **BOLD HEADER** → Neon Green Bebas display
        bold_only = re.match(r"^\*\*([^*]+)\*\*$", raw)
        if bold_only:
            run = p.add_run(bold_only.group(1))
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = NEON_GREEN
            _set_font_all_axes(run, DISPLAY_FONT)
            p.space_before = Pt(5)
            p.space_after  = Pt(3)

        # - bullet item
        elif raw.startswith("- "):
            dot = p.add_run("• ")
            dot.font.size = Pt(10)
            dot.font.color.rgb = NEON_GREEN
            _set_font_all_axes(dot, BODY_FONT)
            _inline_callout_text(p, raw[2:])
            p.space_before = Pt(1)
            p.space_after  = Pt(1)
            p.paragraph_format.left_indent = Inches(0.15)

        # Regular body text
        else:
            _inline_callout_text(p, raw)
            p.space_before = Pt(2)
            p.space_after  = Pt(2)

    post = doc.add_paragraph()
    post.space_before = Pt(0)
    post.space_after  = Pt(10)


def _inline_callout_text(paragraph, text: str) -> None:
    """White body text; **bold** renders Neon Green."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = NEON_GREEN
            _set_font_all_axes(r, BODY_FONT)
        else:
            r = paragraph.add_run(part)
            r.font.size = Pt(10)
            r.font.color.rgb = WHITE
            _set_font_all_axes(r, BODY_FONT)


# ═══════════════════════════════════════════════════════════
# Markdown renderer
# ═══════════════════════════════════════════════════════════

def _render_markdown(doc: Document, content: str) -> None:  # noqa: C901
    lines      = content.split("\n")
    i          = 0
    prev_empty = False
    h2_count   = 0

    # Cover-section state (before [COVER_END])
    in_cover           = True
    cover_h1_done      = False
    cover_subtitle_done = False

    while i < len(lines):
        line = lines[i]

        # ── [COVER_END]: page break, exit cover mode ─────────────────────
        if line.strip() == "[COVER_END]":
            in_cover = False
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            prev_empty = False
            i += 1
            continue

        # ── Blockquote: status skip or callout box ────────────────────────
        if line.strip().startswith("> "):
            raw_content = line.strip()[2:].strip()
            if any(raw_content.startswith(pfx) for pfx in _STATUS_PREFIXES):
                i += 1
                continue
            callout_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                callout_lines.append(lines[i].strip()[2:].strip())
                i += 1
            _add_callout_box(doc, callout_lines)
            prev_empty = False
            continue

        # ── Markdown table ────────────────────────────────────────────────
        if line.strip().startswith("|") and not re.match(r"^\|[-:\s|]+\|$", line.strip()):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            headers, rows = _parse_table_lines(table_lines)
            if headers:
                # Info/metadata table: header row is | | | (two empty cells)
                if len(headers) == 2 and all(h.strip() == "" for h in headers):
                    _add_info_table(doc, rows)
                else:
                    hdr_color = DARK_BLUE if h2_count % 2 == 1 else ELEC_BLUE
                    _add_brand_table(doc, headers, rows, header_color=hdr_color)
            prev_empty = False
            continue

        # ── H1 ────────────────────────────────────────────────────────────
        if line.startswith("# "):
            title = line[2:].strip()
            if in_cover and not cover_h1_done and " & " in title:
                # Cover title: split at " & " into two lines
                cover_h1_done = True
                pre_amp, post_amp = title.split(" & ", 1)

                # Line 1: Electric Blue, 24pt
                p1 = doc.add_paragraph()
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p1.add_run(pre_amp)
                r1.bold = True
                r1.font.size = Pt(24)
                r1.font.color.rgb = ELEC_BLUE
                _set_font_all_axes(r1, DISPLAY_FONT)
                p1.space_before = Pt(12)
                p1.space_after  = Pt(0)

                # Line 2: Dark Blue, 38pt (larger)
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run("& " + post_amp)
                r2.bold = True
                r2.font.size = Pt(38)
                r2.font.color.rgb = DARK_BLUE
                _set_font_all_axes(r2, DISPLAY_FONT)
                p2.space_before = Pt(0)
                p2.space_after  = Pt(10)

            else:
                p = doc.add_paragraph()
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(28)
                r.font.color.rgb = DARK_BLUE
                _set_font_all_axes(r, DISPLAY_FONT)
                p.space_before = Pt(8)
                p.space_after  = Pt(12)
            prev_empty = False

        # ── H2 ────────────────────────────────────────────────────────────
        elif line.startswith("## "):
            h2_count += 1
            in_cover  = False  # any H2 exits cover mode
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.size = Pt(18)
            r.font.color.rgb = DARK_BLUE if h2_count % 2 == 1 else ELEC_BLUE
            _set_font_all_axes(r, DISPLAY_FONT)
            p.space_before = Pt(18)
            p.space_after  = Pt(5)
            prev_empty = False

        # ── H3 ────────────────────────────────────────────────────────────
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = ELEC_BLUE
            _set_font_all_axes(r, DISPLAY_FONT)
            p.space_before = Pt(12)
            p.space_after  = Pt(4)
            prev_empty = False

        # ── Bullet list ───────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            _inline_format(p, line[2:].strip())
            p.space_after = Pt(3)
            prev_empty = False

        # ── Numbered list ─────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            _inline_format(p, re.sub(r"^\d+\.\s", "", line).strip())
            p.space_after = Pt(3)
            prev_empty = False

        # ── Horizontal rule ───────────────────────────────────────────────
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.space_before = Pt(10)
            p.space_after  = Pt(10)
            r = p.add_run("─" * 90)
            r.font.size = Pt(7)
            r.font.color.rgb = ELEC_BLUE
            _set_font_all_axes(r, BODY_FONT)
            prev_empty = False

        # ── Empty line ────────────────────────────────────────────────────
        elif line.strip() == "":
            if not prev_empty:
                sp = doc.add_paragraph()
                sp.space_after = Pt(2)
            prev_empty = True
            i += 1
            continue

        # ── Body paragraph / cover subtitle ──────────────────────────────
        else:
            stripped = line.strip()
            if in_cover and cover_h1_done and not cover_subtitle_done:
                # First plain-text line after cover H1 → italic gray subtitle
                cover_subtitle_done = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(stripped)
                r.italic = True
                r.font.size = Pt(13)
                r.font.color.rgb = MID_GRAY
                _set_font_all_axes(r, BODY_FONT)
                p.space_after = Pt(16)
            else:
                p = doc.add_paragraph()
                _inline_format(p, stripped)
                p.space_after = Pt(6)
            prev_empty = False

        i += 1


def _inline_format(paragraph, text: str) -> None:
    """**bold**, *italic*, plain — 11pt Calibri, all 4 font axes."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(11)
            _set_font_all_axes(r, BODY_FONT)
        else:
            for sub in re.split(r"(\*[^*]+\*)", part):
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    r = paragraph.add_run(sub[1:-1])
                    r.italic = True
                    r.font.size = Pt(11)
                    _set_font_all_axes(r, BODY_FONT)
                else:
                    r = paragraph.add_run(sub)
                    r.font.size = Pt(11)
                    _set_font_all_axes(r, BODY_FONT)
